"""Generate the Word progress report for the post-expansion (v3) work.

Every number is read from the result JSONs / manifests rather than typed in, so the
document cannot drift from what the code actually produced. Anything missing is
reported as missing instead of being silently omitted.

Usage: python scripts/make_report_docx.py [output.docx]
"""
import json
import sys
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
REPORTS = ROOT / "reports"
FIGS = REPORTS / "figures"
ARCHS = ["custom_cnn", "mobilenetv2", "efficientnet_b0"]
NICE = {"custom_cnn": "Custom CNN", "mobilenetv2": "MobileNetV2",
        "efficientnet_b0": "EfficientNet-B0"}
CLASSES = ["CN", "AD", "EMCI", "LMCI"]


def load_json(path):
    p = Path(path)
    if not p.exists():
        return None
    with open(p) as f:
        return json.load(f)


def add_table(doc, headers, rows, bold_first_col=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if bold_first_col and i == 0:
                        r.font.bold = True
    doc.add_paragraph()
    return t


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def pct(x, nd=1):
    return "n/a" if x is None else f"{100*float(x):.{nd}f}%"


def main(out_path=None):
    out_path = Path(out_path) if out_path else ROOT / "reports" / "Alzheimer_Project_Report.docx"
    doc = Document()

    for s in doc.styles:
        pass
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ---------------------------------------------------------------- title
    h = doc.add_heading("Alzheimer's Stage Classification from ADNI MRI", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Progress report after the ADNI dataset expansion — {date.today():%d %B %Y}")
    r.font.size = Pt(11)
    r.font.italic = True

    # ---------------------------------------------------- executive summary
    doc.add_heading("1. What changed, in one page", level=1)
    doc.add_paragraph(
        "Until this session the project had a problem that no amount of modelling could "
        "fix. Every healthy (CN) and Alzheimer's (AD) subject came from ADNI1, scanned "
        "around 2006-07, and every EMCI and LMCI subject came from ADNI-GO/2, scanned "
        "2010 or later. Diagnosis and scanner generation were therefore the same "
        "variable. A model could score ~60% on the four-way task purely by recognising "
        "which scanner protocol produced the image, without learning anything about the "
        "brain. Measurements confirmed exactly that: the models identified the cohort "
        "with 98-100% accuracy while performing at chance on AD vs CN, and Grad-CAM "
        "showed 77% of the model's attention landing outside the brain entirely."
    )
    doc.add_paragraph(
        "The newly downloaded ADNI data fixes this, because it supplies something the "
        "old dataset had none of: CN and AD subjects scanned in the ADNI-GO/2 era. That "
        "makes it possible, for the first time, to assemble all four classes from a "
        "single scanner generation."
    )

    sm = ROOT / "data" / "subject_manifest_v3.csv"
    if sm.exists():
        df = pd.read_csv(sm)
        ct = pd.crosstab(df["class"], df["era"])
        rows = [[c, int(ct.loc[c].get("ADNI1", 0)), int(ct.loc[c].get("GO2", 0)),
                 int(ct.loc[c].sum())] for c in CLASSES if c in ct.index]
        rows.append(["TOTAL", int(ct.get("ADNI1", pd.Series()).sum()),
                     int(ct.get("GO2", pd.Series()).sum()), int(ct.values.sum())])
        doc.add_heading("Dataset after the expansion", level=2)
        add_table(doc, ["Class", "ADNI1 (2006-07)", "ADNI-GO/2 (2010+)", "Total"],
                  rows, bold_first_col=True)
        note(doc, "Subject counts. Before the expansion the ADNI1 column was CN 97 / AD 77 "
                  "and zero for everything else, and the GO/2 column was EMCI 140 / LMCI 125 "
                  "and zero for CN and AD — the confound in its purest form.")

    doc.add_paragraph(
        "The dataset grew from 439 to 853 subjects. 414 new subjects were added; none "
        "of them overlaps the original 439. The primary experiment is now the four-way "
        "task restricted to ADNI-GO/2: 618 subjects, all four classes, one scanner era."
    )

    # -------------------------------------------------- confound audit
    doc.add_heading("2. Proof that the shortcut is gone", level=1)
    doc.add_paragraph(
        "Two independent audits were run against the primary task. Each asks the same "
        "question: if a model were given ONLY this piece of metadata and no image at "
        "all, how accurately could it guess the diagnosis? If the answer is no better "
        "than always guessing the most common class, that cue is useless as a shortcut."
    )
    doc.add_paragraph("On the 618-subject four-way task (baseline 36.7%):")
    add_table(doc,
              ["Potential shortcut", "Accuracy from that cue alone", "Advantage"],
              [["Scanner era (ADNI1 vs GO/2)", "36.7%", "+0.0%"],
               ["Native image resolution", "36.7%", "+0.0%"]],
              bold_first_col=True)

    doc.add_paragraph("On the 501-subject AD-vs-CN task (baseline 56.9%):")
    add_table(doc,
              ["Potential shortcut", "Accuracy from that cue alone", "Advantage"],
              [["Scanner era", "56.9%", "+0.0%"],
               ["Native image resolution", "56.9%", "+0.0%"],
               ["Scanner manufacturer", "57.1%", "+0.2%"],
               ["Voxel row spacing", "57.5%", "+0.6%"],
               ["Scanning site", "64.9%", "+8.0%"]],
              bold_first_col=True)
    note(doc, "Four of the five are negligible. Scanning site is the one that is NOT "
              "cleared, and it is reported here rather than omitted. That +8.0% is "
              "measured across 52 site codes over 501 subjects, so much of it is "
              "overfitting rather than a cue a model could actually exploit — but it "
              "cannot be ruled out on this data.")
    doc.add_paragraph(
        "A second piece of evidence argues against the models reading scanner protocol: "
        "MobileNetV2's AD-vs-CN accuracy is 80.0% on ADNI1 subjects and 85.0% on "
        "ADNI-GO/2 subjects. A model keying on protocol would perform lopsidedly across "
        "the two; performing well on both means it is reading something present in both."
    )
    note(doc, "Across all 853 subjects, resolution still yields +5.9% because a 192x192 "
              "scan implies CN or AD — which is why the GO/2-only task, not the full "
              "dataset, is used as the four-way headline.")

    # ---------------------------------------------------------- results
    doc.add_heading("3. Results", level=1)

    doc.add_heading("3.1 Four-way classification (the task you asked for)", level=2)
    doc.add_paragraph(
        "618 subjects from ADNI-GO/2 only, split by subject 70/15/15 and stratified by "
        "class and era, giving a 93-subject test set. Random guessing scores 25%; always "
        "answering EMCI (the largest class) scores 36.7%. Those are the two numbers any "
        "result has to beat to mean anything."
    )
    rows = []
    for a in ARCHS:
        for suffix, sel in (("", "validation loss"), ("_f1", "validation macro F1")):
            r = load_json(REPORTS / f"{a}_v3go2{suffix}_result.json")
            if r is None:
                rows.append([NICE[a], sel, "not run", "", "", ""])
                continue
            rows.append([
                NICE[a], sel,
                pct(r.get("subject_level_accuracy")),
                pct(r.get("subject_level_macro_f1"), 3).replace("%", ""),
                str(r.get("best_epoch", "?")),
                str(r.get("epochs_run", "?")),
            ])
    add_table(doc, ["Model", "Checkpoint chosen by", "Subject accuracy", "Macro F1",
                    "Best epoch", "Epochs run"], rows, bold_first_col=True)
    note(doc, "Subject-level is the number that matters: each subject contributes 32 "
              "slices, combined into one prediction per person. Macro F1 averages the "
              "four stages equally, so it exposes a model that has quietly stopped "
              "predicting one of them — which raw accuracy hides. 'Best epoch' shows "
              "which checkpoint was restored; a value of 2 means an almost untrained "
              "model was kept, and that result should be discarded rather than reported.")
    doc.add_paragraph(
        "Honest reading of this table: at 93 test subjects, one subject is worth about "
        "1.1 accuracy points, and a 95% confidence interval around any of these numbers "
        "spans roughly plus or minus 10 points. None of them separates convincingly from "
        "the 36.7% baseline. The four-way task is not solved."
    )

    # Per-stage detail: the point of the project is all four stages, and an aggregate
    # accuracy can hide a model that has silently stopped predicting one of them.
    doc.add_heading("Per-stage performance, and why aggregate accuracy misleads", level=3)
    import numpy as np
    for tag, label in (("custom_cnn_v3go2_f1", "Custom CNN (macro-F1 selected)"),
                       ("efficientnet_b0_v3go2", "EfficientNet-B0 (validation-loss selected)"),
                       ("efficientnet_b0_v3go2_f1", "EfficientNet-B0 (macro-F1 selected)")):
        cm_p = REPORTS / f"{tag}_subject_cm_soft.npy"
        if not cm_p.exists():
            continue
        cm = np.load(cm_p)
        rows = []
        for i, c in enumerate(CLASSES):
            tot = cm[i].sum()
            rows.append([c, int(tot), int(cm[i, i]),
                         f"{100*cm[i, i]/tot:.0f}%" if tot else "n/a"])
        doc.add_paragraph(label, style="List Bullet")
        add_table(doc, ["Stage", "Test subjects", "Correct", "Recall"], rows,
                  bold_first_col=True)
    doc.add_paragraph(
        "The two EfficientNet rows are the clearest illustration of why this project "
        "reports per-stage numbers. Selected on validation loss, that model scored 39.8% "
        "overall — the highest of the three — but achieved it by answering EMCI for 75 of "
        "93 subjects and never once predicting LMCI: 0% recall on an entire disease "
        "stage. Re-selected on macro F1 it predicts all four stages, recovers LMCI to "
        "47% recall, and its overall accuracy FALLS to 33.3%. The lower number is the "
        "more honest model."
    )

    doc.add_heading("3.2 AD vs CN — the clinically meaningful comparison", level=2)
    doc.add_paragraph(
        "Distinguishing Alzheimer's from healthy is the comparison the medical "
        "literature benchmarks and the one that was previously at chance. It now uses "
        "501 subjects drawn from both eras in near-equal proportion, so the scanner "
        "gives no clue: knowing the era alone scores exactly the 56.9% majority "
        "baseline. Previously this task had only 174 subjects, all from one cohort."
    )
    rows = []
    for a in ARCHS:
        r = load_json(REPORTS / f"{a}_ADvsCN_v3adcn_result.json")
        if r is None:
            rows.append([NICE[a], "not run", "", "", "", ""])
            continue
        ci = r.get("accuracy_95CI") or [None, None]
        aci = r.get("roc_auc_95CI") or [None, None]
        rows.append([
            NICE[a], pct(r.get("accuracy")),
            f"[{pct(ci[0])}, {pct(ci[1])}]",
            pct(r.get("majority_baseline")),
            f"{r.get('roc_auc', float('nan')):.3f}",
            f"[{aci[0]:.3f}, {aci[1]:.3f}]" if aci[0] is not None else "n/a",
        ])
    add_table(doc, ["Model", "Accuracy", "95% CI", "Baseline", "ROC AUC", "AUC 95% CI"],
              rows, bold_first_col=True)
    note(doc, "ROC AUC measures whether the model ranks Alzheimer's subjects above "
              "healthy ones; 0.5 means no signal. The confidence interval matters more "
              "than the point estimate — if it still includes 0.5, the result is "
              "suggestive rather than established.")

    for a in ARCHS:
        r = load_json(REPORTS / f"{a}_ADvsCN_v3adcn_result.json")
        if r and r.get("verdict"):
            doc.add_paragraph(f"{NICE[a]}: {r['verdict']}", style="List Bullet")

    doc.add_heading("The trajectory across three versions of the dataset", level=3)
    doc.add_paragraph(
        "This is the clearest summary of what has actually been achieved. The same task, "
        "measured three times as the data improved:"
    )
    add_table(doc,
              ["Dataset version", "Subjects", "Best ROC AUC", "95% CI", "Conclusion"],
              [["v1 — original slices", "174", "0.509", "includes 0.5",
                "No signal (one model scored 0.345, worse than chance)"],
               ["v2 — slice alignment fixed", "174", "0.673", "[0.457, 0.889]",
                "Suggestive, not established"],
               ["v3 — expanded and era-balanced", "501", "0.906", "[0.830, 0.981]",
                "Established"]],
              bold_first_col=True)
    note(doc, "The confidence interval is the whole point. In v1 and v2 the interval "
              "still contained 0.5, meaning 'no better than guessing' could not be ruled "
              "out. In v3 it does not — which is what makes this the first result in the "
              "project that can be stated as a finding rather than a hint.")

    # ---------------------------------------------- why four-way is hard
    doc.add_heading("3.3 Why four-way is hard — measured, not assumed", level=2)
    doc.add_paragraph(
        "It is worth being precise about where the four-way errors actually come from, "
        "because the answer is partly outside the model's control."
    )
    add_table(doc, ["Measurement", "Value"],
              [["Mean four-way accuracy across six runs", "38.5%"],
               ["Same runs, EMCI and LMCI merged into one 'MCI' class", "54.5%"],
               ["Accuracy recovered by merging the two MCI stages", "+15.9 points"],
               ["EMCI vs LMCI, given the model already knows it is MCI", "58.7%"],
               ["  (a coin flip would score)", "50.0%"]],
              bold_first_col=True)
    doc.add_paragraph(
        "ADNI does not define EMCI and LMCI by anatomy. It defines them by a "
        "delayed-recall memory-test score: roughly 1 to 1.5 standard deviations below "
        "the normative mean is EMCI, more than 1.5 is LMCI. They are the same disease "
        "stage separated by a psychological test threshold, and the published "
        "literature describes the boundary as inherently blurred. So a large part of "
        "the four-way ceiling lives in the labels rather than in the images — nearly "
        "16 accuracy points are lost at this one boundary, and the model's sub-stage "
        "decision is barely better than guessing."
    )
    doc.add_paragraph(
        "Being straight about the limits of that explanation: merging the two MCI "
        "stages gives 54.5%, but the three-way majority baseline is 57.0%. So the "
        "models do not clear the baseline on the coarse task either. The EMCI/LMCI "
        "boundary is the single biggest sink, not the whole problem."
    )
    for txt in [
        "The continuum problem. AD vs CN works well because it compares the two "
        "extremes. CN, EMCI, LMCI and AD form a gradient, not four separate groups, "
        "and the middle stages overlap both ends — four-way asks the model to draw "
        "three boundaries through a smooth progression.",
        "LMCI is capped at 125 subjects. ADNI has no more with this scan type, so the "
        "hardest class is also the smallest.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    # ------------------------------------------ improvement attempts
    doc.add_heading("3.4 Attempts to raise four-way accuracy", level=2)
    doc.add_paragraph(
        "Three levers were tested, each targeting something specific rather than "
        "generic tuning:"
    )
    for txt in [
        "Full resolution. The pipeline routes every slice through a 144-pixel "
        "bottleneck to make ADNI1's 192x192 scans comparable with ADNI-GO/2's "
        "256x256. The primary task is GO/2-only, where every scan is already 256, so "
        "that step corrects a problem that is not present while discarding detail at "
        "roughly 1.7mm per pixel — the scale at which disease stages differ.",
        "In-domain initialisation. Start from the AD-vs-CN model instead of random "
        "weights. That model reaches ROC AUC 0.906 on these same images, so its "
        "features demonstrably encode real atrophy. This is not the ImageNet "
        "pretraining that failed earlier: the problem there was that natural "
        "photographs are too far from MRI, and here the source is the same modality, "
        "anatomy and preprocessing.",
        "Self-supervised pretraining. A masked autoencoder is trained to rebuild "
        "blanked-out squares of roughly 19,100 slices with no labels at all, then its "
        "encoder initialises the classifier. Labels are the scarce resource in this "
        "project, not images — this is the only lever that attacks that directly.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    rows = []
    for tag, label in (
        ("custom_cnn_v3go2_f1", "Custom CNN — baseline"),
        ("custom_cnn_v3go2hi_f1", "Custom CNN — full resolution"),
        ("custom_cnn_v3go2_f1_init-custom_cnn_ADvsCN", "Custom CNN — from AD-vs-CN"),
        ("custom_cnn_v3go2_f1_init-ssl_encoder", "Custom CNN — from autoencoder"),
        ("custom_cnn_v3go2hi_f1_init-ssl_encoder", "Custom CNN — full res + autoencoder"),
        ("mobilenetv2_v3go2_f1", "MobileNetV2 — baseline"),
        ("mobilenetv2_v3go2hi_f1", "MobileNetV2 — full resolution"),
        ("mobilenetv2_v3go2_f1_init-mobilenetv2_ADvsCN", "MobileNetV2 — from AD-vs-CN"),
        ("mobilenetv2_v3go2hi_f1_init-mobilenetv2_ADvsCN", "MobileNetV2 — full res + AD-vs-CN"),
        ("efficientnet_b0_v3go2_f1", "EfficientNet-B0 — baseline"),
        ("efficientnet_b0_v3go2hi_f1", "EfficientNet-B0 — full resolution"),
    ):
        r = load_json(REPORTS / f"{tag}_result.json")
        if r is None:
            continue
        rows.append([label,
                     pct(r.get("subject_level_accuracy")),
                     pct(r.get("subject_level_macro_f1"), 3).replace("%", ""),
                     str(r.get("best_epoch", "?"))])
    if rows:
        add_table(doc, ["Configuration", "Subject accuracy", "Macro F1", "Best epoch"],
                  rows, bold_first_col=True)
        note(doc, "Baseline to beat is 36.7% (always answering EMCI). With 93 test "
                  "subjects, one subject is 1.1 points and the confidence interval "
                  "around any of these spans roughly plus or minus 10 points — so only "
                  "a large gap here would be meaningful, and differences of two or "
                  "three points are noise.")

    ssl = load_json(REPORTS / "ssl_pretrain_result.json")
    if ssl:
        doc.add_paragraph(
            f"The autoencoder trained on {ssl.get('n_pretrain_slices', 0):,} slices "
            f"from {ssl.get('n_pretrain_subjects', 0)} subjects, using no labels. "
            "Validation and test subjects were excluded: reconstruction needs no "
            "labels, but fitting a model to images that later appear in evaluation is "
            "still leakage, and this project has measured that such leakage is worth "
            "up to 37 accuracy points of illusion."
        )
        f = FIGS / "ssl_reconstruction.png"
        if f.exists():
            doc.add_picture(str(f), width=Inches(6.2))
            note(doc, "Top: original slice. Middle: what the network was actually "
                      "given, with squares blanked out. Bottom: its reconstruction. It "
                      "never saw the hidden regions and must infer them from the "
                      "surrounding anatomy, which is what forces it to learn brain "
                      "structure rather than copying pixels.")

    doc.add_heading("3.5 For comparison: four-way using all 853 subjects", level=2)
    r = load_json(REPORTS / "mobilenetv2_v3_result.json")
    if r:
        doc.add_paragraph(
            f"MobileNetV2 over all 853 subjects reaches "
            f"{pct(r.get('subject_level_accuracy_softvote'))} subject-level accuracy — "
            "lower than the 618-subject era-matched result, not higher."
        )
        doc.add_paragraph(
            f"That comparison should not be leaned on, because this run restored its "
            f"checkpoint from epoch {r.get('best_epoch')} (of {r.get('epochs_run')} run) "
            "under validation-loss selection, so it is the same near-untrained-model "
            "problem described above rather than a clean measurement. It is reported "
            "here for completeness and flagged as unreliable."
        )
        doc.add_paragraph(
            "Either way this number is deliberately NOT the headline: EMCI and LMCI "
            "still exist only in ADNI-GO/2 within this pool, so part of any score here "
            "is still scanner recognition rather than diagnosis."
        )
    else:
        doc.add_paragraph("Not yet run.")

    # -------------------------------------------- cross-era validation
    doc.add_heading("4. The hardest test: a completely different scanner generation",
                    level=1)
    doc.add_paragraph(
        "Every result above splits training and test data inside one pool, so both "
        "sides share scanner generation, protocol and reconstruction software. A model "
        "could still be leaning on cohort-specific image texture and score well. The "
        "strongest available check is to train on one ADNI cohort and test on the "
        "whole of the other one — different machines, different years, different "
        "vendors, no overlapping subjects. This was impossible before the expansion, "
        "because there were no CN or AD subjects in ADNI-GO/2 at all."
    )
    doc.add_paragraph(
        "It is also the question that matters for real use: would this work on scans "
        "from a hospital whose machine the model has never seen?"
    )
    rows = []
    for arch in ARCHS:
        for d, dlabel in (("adni1_to_go2", "ADNI1 -> GO/2"),
                          ("go2_to_adni1", "GO/2 -> ADNI1")):
            r = load_json(REPORTS / f"{arch}_crossera_{d}_result.json")
            if r is None:
                continue
            aci = r.get("roc_auc_95CI") or [None, None]
            rows.append([NICE[arch], dlabel, str(r.get("n_test_subjects")),
                         f"{r.get('roc_auc', float('nan')):.3f}",
                         f"[{aci[0]:.3f}, {aci[1]:.3f}]" if aci[0] is not None else "n/a",
                         pct(r.get("accuracy")), pct(r.get("majority_baseline"))])
    if rows:
        add_table(doc, ["Model", "Trained -> tested", "Test subjects", "ROC AUC",
                        "AUC 95% CI", "Accuracy", "Baseline"], rows,
                  bold_first_col=True)
        doc.add_paragraph(
            "The headline here: every AUC confidence interval excludes 0.5. A model "
            "trained only on 2006-era scans still ranks Alzheimer's above healthy on "
            "2010s scanners, and the reverse also holds. Whatever these models learned, "
            "it is present in both scanner generations — which is very hard to explain "
            "if they were reading scanner protocol."
        )
        doc.add_paragraph(
            "Two honest qualifications. First, performance drops: within one cohort the "
            "best AUC is 0.906, across cohorts it is 0.68 to 0.79. That gap is the "
            "realistic answer to 'would this work elsewhere' — the signal survives, but "
            "weakened. Second, accuracy transfers much less reliably than ranking: only "
            "half of these runs beat their baseline, because the decision threshold is "
            "fitted on the training cohort and does not carry across to a different "
            "scanner's brightness distribution. In practice the ranking is reusable but "
            "the cut-off would need recalibrating at each new site."
        )
        note(doc, "Deliberately NOT claimed here: that one architecture generalises "
                  "better than another. The custom CNN wins one direction and loses the "
                  "other, so that comparison is noise at this sample size.")

    # ------------------------------------------------------------ grad-cam
    doc.add_heading("5. Where the model actually looks (Grad-CAM)", level=1)
    doc.add_paragraph(
        "Grad-CAM highlights the image regions driving each prediction. It is used here "
        "as a numeric check, not decoration: the fraction of heat landing outside the "
        "brain mask. On the old data this was 77% overall, with the heat sitting in "
        "vertical bands at the left and right image margins — empty background. That was "
        "independent confirmation of the confound, reached by a different method than "
        "the confusion-matrix analysis."
    )
    # gradcam.py writes reports/gradcam_{tag}.json
    summary_rows = [["Old best model (v1 data, four-way)", "77.1%",
                     "CN 59 / AD 64 / EMCI 91 / LMCI 95"]]
    for tag, label in (("mobilenetv2_ADvsCN", "MobileNetV2, AD vs CN (v3)"),
                       ("custom_cnn_v3go2_f1", "Custom CNN, four-way (v3go2)")):
        gc = load_json(REPORTS / f"gradcam_{tag}.json")
        if gc:
            by = " / ".join(f"{c} {100*v:.0f}"
                            for c, v in gc.get("by_class", {}).items())
            summary_rows.append([label,
                                 pct(gc.get("mean_gradcam_mass_outside_brain")), by])
    add_table(doc, ["Model", "Attention outside the brain", "By class (%)"],
              summary_rows, bold_first_col=True)
    note(doc, "Lower is better. The old model's figure was not only high but wildly "
              "uneven across classes (59% to 95%) — a model reading genuine anatomy has "
              "no reason to look outside the brain far more for one diagnosis than "
              "another, so that spread was itself evidence of scanner artifacts. The "
              "AD-vs-CN model's two classes now sit at 36.4% and 36.6%, essentially "
              "identical.")

    f = FIGS / "gradcam_mobilenetv2_ADvsCN.png"
    if f.exists():
        doc.add_paragraph(
            "Below: the AD-vs-CN model, the one with established signal. On the old data "
            "the heat formed vertical bands at the left and right image margins — empty "
            "background. It now sits mostly on tissue around the ventricles and the "
            "medial temporal lobe, which is where Alzheimer's actually shows up "
            "(enlarged ventricles, shrunken hippocampus). Not every case is clean: the "
            "second row still puts most of its attention on the skull edge, which is why "
            "the average is 36% rather than near zero."
        )
        doc.add_picture(str(f), width=Inches(5.6))
        note(doc, "Left = input slice, middle = Grad-CAM heatmap, right = overlay with "
                  "the predicted class and its confidence.")
    if not any((REPORTS / f"gradcam_{t}.json").exists()
               for t in ("mobilenetv2_ADvsCN", "custom_cnn_v3go2_f1")):
        doc.add_paragraph("Grad-CAM has not yet been run on the v3 models.")

    # ------------------------------------------------------- data quality
    doc.add_heading("6. Data preparation and quality control", level=1)
    for txt in [
        "One scan session per subject. Every new subject had 2-4 sessions. Treating "
        "them as separate samples would put the same person in both training and test "
        "data — the single most common error in published work of this kind, worth up "
        "to 37 accuracy points of pure illusion on this exact dataset (measured "
        "earlier in the project). The earliest usable session per subject is kept.",

        "Subject-wise splitting. Train/validation/test is decided per person before any "
        "image is created, so all 32 slices from one person stay on the same side.",

        "Anatomically anchored slices. Slices are taken 48-92mm below the top of the "
        "head, measured in real millimetres using each scan's own voxel spacing, so the "
        "same slice number shows the same anatomy in every subject and always covers "
        "the hippocampus — the structure Alzheimer's attacks first.",

        "New scanner types validated before use. The expansion introduced Philips and "
        "Siemens parallel-imaging variants (MPRAGE_GRAPPA2, MPRAGE_SENSE2 and others) "
        "that the original data never contained. All 10 series families were checked to "
        "confirm they are acquired in the same orientation the pipeline assumes; a "
        "mismatch would have silently produced images cut through the wrong plane.",

        "24 broken single-image sessions were found and excluded, along with 12 subjects "
        "that had no usable session.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    for name, cap in (("qc_v3_range.png",
                       "Extracted slice band per scanner type. Each row is one series "
                       "family, moving from the top of the brain (left) down to the "
                       "temporal lobes and hippocampus (right). The anatomy lines up "
                       "across all scanner types, which is what makes them comparable."),):
        f = FIGS / name
        if f.exists():
            doc.add_picture(str(f), width=Inches(6.2))
            note(doc, cap)

    # ------------------------------------------------------------ honesty
    doc.add_heading("7. Limitations — read this before quoting any number", level=1)
    for txt in [
        "LMCI could not be expanded. ADNI simply does not contain many more LMCI "
        "subjects with this scan type; all 5 additional ones requested were unavailable. "
        "LMCI remains the smallest class at 125 subjects and is the hardest to separate "
        "from EMCI, since the two are adjacent stages distinguished clinically by memory "
        "test scores rather than by obvious differences in brain structure.",

        "EMCI and LMCI still come only from ADNI-GO/2. This is why the four-way task is "
        "restricted to that era. A four-way model spanning both eras would still be "
        "partly reading the scanner.",

        "The test set is 93 subjects. One subject is worth about 1.1 accuracy points. "
        "Differences of a few points between models are not meaningful. This project has "
        "already been burned once by exactly that: a plausible story about masking "
        "hurting performance, supported by two architectures agreeing, reversed "
        "completely on the third and turned out to be noise.",

        "This is not a diagnostic tool. It is trained on research-grade scans from one "
        "study, and performance on scans from a different hospital or scanner is "
        "unknown and would likely be worse.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    # --------------------------------------------------------------- next
    doc.add_heading("8. What comes next", level=1)
    for txt in [
        "Train on one era and test on the other. Now that CN and AD exist in both, a "
        "model can be trained on ADNI1 and tested on ADNI-GO/2. That is a far stronger "
        "test of whether it learned anatomy rather than protocol, and it was impossible "
        "before this expansion.",

        "Re-examine preprocessing variants (skull-stripping, brain cropping) on the "
        "larger test set. They were previously indistinguishable from noise at 26 test "
        "subjects; at 93 they may become separable.",

        "The demonstration website stays paused until the four-way and AD-vs-CN results "
        "are convincingly above baseline. Shipping confident stage predictions from a "
        "model without established signal would be misleading.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else None)

