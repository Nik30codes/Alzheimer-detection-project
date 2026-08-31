"""Builds a SIMPLIFIED-LANGUAGE version of the ETE research paper -- same real data,
same tables, same figures as build_report.py, but the prose is rewritten in plainer,
more first-person, less "academic-sounding" language. Separate output file; does not
touch build_report.py's output or anything the user edited by hand.
Run: python report_assets/build_report_simple.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Nikunj_Bhalla_2430030052_ETE_Research_Paper_Simple.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(htext)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")
    return t


def figure(path, caption, width=4.6):
    if not os.path.exists(path):
        para(f"[figure missing: {os.path.basename(path)}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)


FIG = os.path.join(HERE)
DEMO_FIG = os.path.join(HERE, "..", "demo", "results")

# ---------------------------------------------------------------- TITLE ----
title = doc.add_heading(
    "Classifying Alzheimer's Disease from Real ADNI MRI Scans, and Checking "
    "Whether the Model Is Actually Cheating", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Nikunj Bhalla")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -------------------------------------------------------------- ABSTRACT ---
h1("ABSTRACT")
para(
    "Being able to spot Alzheimer's disease automatically from an MRI scan is "
    "genuinely useful, since the brain changes it causes show up before the disease "
    "gets bad. The problem is that a lot of papers in this area report accuracy in "
    "the 80-99% range, and a good chunk of that comes from test setups that "
    "accidentally let information about a person or a scanner leak from training "
    "into testing. In this project I built the classifier from raw ADNI DICOM "
    "scans -- 853 people across two different scanner generations, not a ready-made "
    "image dataset -- and before trusting any result, I went looking for four "
    "specific ways the pipeline could be cheating: the same person's slices ending "
    "up on both sides of the split, the scanner generation lining up with the "
    "diagnosis, the slices being cut from the wrong part of the brain, and a "
    "geometry problem from how the scans get reconstructed. I fixed three of these. "
    "The fourth, which hospital the scan came from, I measured but haven't fixed "
    "yet, so I'm reporting it honestly instead of hiding it. I compared three "
    "models -- a small CNN I built myself, MobileNetV2, and EfficientNet-B0 -- all "
    "trained from scratch, because starting them from ImageNet weights actually "
    "made them worse on this data. For telling Alzheimer's apart from cognitively "
    "normal scans, 5-fold cross-validation across 501 subjects gives 74.1% accuracy "
    "and an ROC AUC of 0.784 (95% CI [0.743, 0.826]), which lines up with what "
    "properly-split studies in the literature get, and with a recent large MRI "
    "foundation model tested the same way. I also tried the harder four-stage "
    "version of the task (normal / early MCI / late MCI / Alzheimer's), but that "
    "result isn't solid yet -- it's still tangled up with one of the leftover "
    "geometry issues, so I'm reporting it as unresolved rather than pretending it's "
    "done."
)

# ----------------------------------------------------------- INTRODUCTION --
h1("1. INTRODUCTION")
para(
    "Alzheimer's disease is the most common cause of dementia. It's a slow loss of "
    "memory and independence that gets worse over years, and MRI matters here "
    "because you can actually see shrinkage in the hippocampus and nearby brain "
    "regions before someone reaches the late stages. If a model could read that "
    "signal reliably, it could help catch things earlier and more consistently "
    "than relying on a radiologist reading every scan by eye."
)
para(
    "The hard part isn't training a CNN that scores well on a test set. It's "
    "making sure it actually learned to recognize atrophy instead of learning a "
    "shortcut. Two things about MRI datasets make shortcuts easy to fall into "
    "without noticing. First, each person contributes dozens of slices that look "
    "almost identical to each other, so if a split isn't done person-by-person, the "
    "model can partly just memorize individual people instead of learning anything "
    "general. Second, things like which scanner, which coil, and what resolution "
    "was used often line up with the diagnosis for reasons that have nothing to do "
    "with the brain -- usually because different groups of patients were scanned "
    "years apart on different machines. A model can score well by picking up on "
    "either of these instead of the actual disease."
)
para(
    "I treated that as the main problem to solve, not a footnote. Starting from "
    "853 people's worth of raw ADNI DICOM files, the pipeline rebuilds a 3D volume "
    "for each person, reslices it the standard way, and pulls out a slice band "
    "anchored in millimetres rather than a rough fraction of the image, before any "
    "train/test split even happens. Then I went looking for four specific ways "
    "this could still be cheating, tested each one directly, and fixed the three "
    "that turned out to be real problems. What this paper actually adds:"
)
bullet("A slice-extraction pipeline that rebuilds the full 3D brain volume from "
       "raw sagittal DICOM files and anchors the slice band in real millimetres, "
       "instead of working from pre-cut images.")
bullet("A direct test for four separate ways the model could be taking shortcuts "
       "-- subject leakage, scanner-era bias, misaligned slices, and a geometry "
       "artifact -- each checked with its own measurement instead of just assumed "
       "fine.")
bullet("A side-by-side comparison of a custom CNN, MobileNetV2, and "
       "EfficientNet-B0, all trained from scratch, plus the finding that starting "
       "from ImageNet weights actually hurts on grayscale MRI.")
bullet("A result that's backed by 5-fold cross-validation with confidence "
       "intervals, not just one lucky train/test split -- because in this exact "
       "pipeline, a single split overstated accuracy by more than ten points.")
bullet("An honest note on what's still unresolved, including a four-stage result "
       "that I'm not claiming is finished.")

# -------------------------------------------------------- LITERATURE REVIEW-
h1("2. LITERATURE REVIEW")
para(
    "CNNs are the go-to tool for MRI-based Alzheimer's classification now, mainly "
    "because they can pull out useful features straight from the pixels instead of "
    "needing someone to hand-design measurements first. A lot of papers use "
    "ImageNet-pretrained backbones like EfficientNet [3] and MobileNet [4], "
    "assuming that low-level features learned from photos still help when there "
    "isn't much medical data to train on. As I show in Section 4, that assumption "
    "didn't hold up here."
)
para(
    "There's a second, smaller thread in the literature that's more about how "
    "these models get evaluated than about architecture. Ansart et al. [8] found "
    "that models trained and tested within one site or cohort often don't hold up "
    "when tested on outside data, which points to cohort-specific shortcuts rather "
    "than real anatomical signal. Tinauer et al. [9] showed that skull-stripping, "
    "usually treated as a harmless preprocessing step, can actually introduce its "
    "own shortcut depending on the dataset -- basically a Clever Hans effect, "
    "which is a good argument for testing preprocessing choices instead of "
    "assuming they're safe. A more recent paper on leakage-aware 3D CNNs [10] "
    "makes subject-level splitting an explicit, stated requirement, which says "
    "something about how often it gets left out elsewhere. And BrainDINO [11], a "
    "foundation model pretrained on 6.6 million MRI slices from twenty different "
    "datasets, reports an AUC of 0.850 (95% CI [0.754, 0.947]) on subject-disjoint "
    "AD-vs-CN classification -- a useful, carefully done reference point for what "
    "this task's ceiling looks like without any shortcut inflation."
)
para(
    "Scanner and site harmonization is its own named problem in neuroimaging more "
    "broadly. Tools like ComBat, and newer versions of it like CovBat and "
    "DeepComBat, exist specifically to strip out scanner- and site-related "
    "variance from imaging features before modelling. My own site-effect "
    "measurement in Section 4.1 lines up with why that literature exists in the "
    "first place; actually applying one of those methods here is left for future "
    "work."
)
para(
    "This project sits somewhere between those two threads: fairly standard CNN "
    "architectures, but evaluated with the kind of care the second thread argues "
    "for -- checking for confounds directly, splitting by subject, and reporting "
    "cross-validated numbers instead of trusting one test accuracy at face value."
)

# ------------------------------------------------------------- METHODOLOGY -
h1("3. METHODOLOGY")

h2("3.1 Dataset Description")
para(
    "I pulled the data directly from the ADNI (Alzheimer's Disease Neuroimaging "
    "Initiative) archive as raw DICOM files, downloaded in two batches that cover "
    "ADNI's earlier phase (roughly 2005-2007) and its later phase (2011 onward) -- "
    "I'll call these ADNI1 and ADNI-GO/2 below. Together they add up to 853 "
    "subjects across four diagnostic groups (Table 1):"
)
table(["Class", "ADNI1", "ADNI-GO/2", "Total"],
      [["Cognitively Normal (CN)", 127, 158, 285],
       ["Alzheimer's Disease (AD)", 108, 108, 216],
       ["Early MCI (EMCI)", 0, 227, 227],
       ["Late MCI (LMCI)", 0, 125, 125]])
para(
    "Each person contributes one T1-weighted MPRAGE-type scan. I split subjects "
    "70/15/15 into train/validation/test at the person level, before generating "
    "a single slice image, so nobody's scans end up on both sides of a split."
)
figure(os.path.join(FIG, "class_distribution.png"),
      "Figure 1: Class distribution across the dataset.")

h2("3.2 Data Preprocessing")
para(
    "For each person, I stack their sagittal DICOM series into a 3D volume and "
    "reslice it into the usual top-down axial view. From there I pull out a "
    "32-slice band, but instead of anchoring it as a rough fraction of image "
    "height, I anchor it a fixed physical distance (48-92mm) below the top of the "
    "skull, using that scan's own pixel spacing. I switched to this after finding "
    "that the fraction-based version drifted onto different anatomy in different "
    "people -- in some cases it missed the hippocampus completely, which is the "
    "single most important structure for this task. Stacking slices that were "
    "each scanned separately also introduces some grain, which I clean up with "
    "Non-Local Means denoising. And since ADNI1 scans are natively lower "
    "resolution than ADNI-GO/2 scans, I route every image through the same "
    "intermediate resolution before the final resize to 224x224, so resolution "
    "itself doesn't quietly become a stand-in for scanner generation -- and, given "
    "the class table above, a stand-in for diagnosis."
)
figure(os.path.join(FIG, "qc_v3_range.png"),
      "Figure 2: Example extracted axial slice range across the anchored band, "
      "showing consistent anatomical coverage.")

h2("3.3 How I Checked for Confounds")
para(
    "Before trusting any accuracy number, I tested each candidate confound the "
    "same way: train a simple model on nothing but the non-image metadata (things "
    "like scanner era, native resolution, voxel spacing) and see if it can guess "
    "the diagnosis better than just picking the majority class -- with zero image "
    "pixels involved. If it can, that variable is a usable shortcut, and any "
    "image-based model might be leaning on it instead of the brain. I ran four of "
    "these tests:"
)
table(["#", "Possible confound", "What I found", "Status"],
      [["1", "Same person's slices split across train and test",
        "+36.9 accuracy points when I deliberately let this happen", "Fixed: split by subject"],
       ["2", "Scanner era (ADNI1 vs GO/2) lining up with class",
        "Era guessable at 95-100% before the fix; diagnosis near chance", "Fixed: balanced the dataset by era"],
       ["3", "Slice band landing on the wrong anatomy",
        "Band drifted per subject; sometimes missed the hippocampus", "Fixed: anchored in millimetres"],
       ["4", "Geometry / stretch from how scans get reconstructed",
        "Metadata alone: +4.2% on the 4-way task; -4.6% (clean) on AD-vs-CN", "Partly fixed (isotropic resampling)"]])
para(
    "There's a fifth one I found too: which imaging site a scan came from. Site "
    "identity alone was worth +8.0% over baseline on the AD-vs-CN task, which "
    "isn't nothing. That number is probably inflated, since there are a lot of "
    "sites relative to how many subjects I have, and the final model's accuracy "
    "is actually pretty balanced across both eras (which argues against it just "
    "reading protocol). Still, I haven't fixed this one, so I'm reporting it as "
    "an open limitation rather than something resolved."
)

h2("3.4 The Models")
para(
    "I compared three architectures, all working on one 224x224 axial slice at "
    "a time:"
)
bullet("A custom CNN (about 1.2M parameters) I built myself: four "
       "conv-batchnorm-ReLU blocks with max-pooling, channel widths going "
       "32-64-128-256, then global average pooling and a small dropout-regularized "
       "head. I used global average pooling instead of flattening straight into a "
       "dense layer specifically to keep the parameter count down given how much "
       "data I actually have.")
bullet("MobileNetV2 (about 2.2M parameters), the standard torchvision version, "
       "using inverted residual blocks with depthwise-separable convolutions.")
bullet("EfficientNet-B0 (about 4.0M parameters), also the torchvision version, "
       "using compound depth/width/resolution scaling.")
para(
    "I also tried starting MobileNetV2 from ImageNet weights, under three "
    "different fine-tuning setups (fully unfrozen, partially unfrozen with "
    "BatchNorm frozen, partially unfrozen with BatchNorm adapting). Every one of "
    "them did worse than the same model trained from scratch, by 15-20 accuracy "
    "points. That makes sense once you think about it -- the pretrained features "
    "are tuned for natural photos, and grayscale MRI copied into three fake color "
    "channels doesn't look much like a photo at all. So every result below uses "
    "random initialization, not pretraining."
)

h2("3.5 Training Setup")
table(["Parameter", "Value"],
      [["Input size", "224 x 224, 1 or 3 channels"],
       ["Optimizer", "Adam, lr 1e-3"],
       ["Weight decay", "1e-4"],
       ["Batch size", "32"],
       ["Loss", "Class-weighted cross-entropy"],
       ["LR schedule", "ReduceLROnPlateau (factor 0.5, patience 2)"],
       ["Early stopping", "Patience 7 epochs, best checkpoint restored"],
       ["Max epochs", "40"],
       ["Augmentation", "Random affine (rotate/translate/scale) + horizontal flip"],
       ["Precision", "Mixed precision (fp16 autocast + gradient scaling)"]])

h2("3.6 How I Evaluated Everything")
para(
    "For each person, I average the softmax probabilities across all their "
    "slices to get one prediction per person, since what actually matters "
    "clinically is the person, not any one 2D image. For the main AD-vs-CN task, "
    "I use 5-fold cross-validation at the subject level: every person gets "
    "exactly one out-of-fold prediction, and I report the pooled result with a "
    "95% Wilson interval for accuracy and a Hanley-McNeil interval for ROC AUC. I "
    "started doing it this way after noticing that one 75-subject test split of "
    "this exact setup gave 82.7% accuracy and AUC 0.906, while 5-fold "
    "cross-validation over all 501 subjects gave 74.1% and AUC 0.784, with "
    "confidence intervals that don't even overlap. In other words, that first "
    "split was just lucky, not a number I should trust. The cutoff for turning a "
    "probability into a label is picked on the validation subjects using Youden's "
    "J statistic, and never touches the test fold."
)

# ------------------------------------------------------ RESULTS/DISCUSSION -
h1("4. RESULTS AND DISCUSSION")

h2("4.1 AD vs Cognitively Normal -- the Main Result")
para(
    "Table 3 shows each model's performance on one held-out test split (75 "
    "people) just for illustration, next to the actual cross-validated result. "
    "Only that cross-validated number should be treated as the real headline, for "
    "the reason explained in Section 3.6."
)
table(["Model", "Split", "n", "Accuracy", "Macro F1", "ROC AUC"],
      [["Custom CNN", "single split", 75, "73.3%", "0.725", "0.838"],
       ["MobileNetV2", "single split", 75, "76.0%", "0.758", "0.852"],
       ["EfficientNet-B0", "single split", 75, "78.7%", "0.786", "0.876"],
       ["MobileNetV2", "5-fold CV (headline)", 501, "74.1% [70.0, 77.7]", "--",
        "0.784 [0.743, 0.826]"]])
para(
    "The cross-validated AUC interval stays well clear of 0.5, so there's a real "
    "ranking signal here, not noise. It also matches the wider literature: "
    "properly subject-split studies on ADNI-sized data tend to report binary "
    "AD-vs-CN accuracy under 71%, and BrainDINO, a foundation model pretrained on "
    "6.6 million MRI slices, reports AUC 0.850 [0.754, 0.947] [11] on the same "
    "comparison -- an interval that overlaps mine, despite BrainDINO using two "
    "orders of magnitude more pretraining data than I have here."
)
para(
    "Below are the confusion matrix and training curves for each of the three "
    "models, all from the same 75-subject test split as Table 3, so you can "
    "compare the error patterns and learning curves directly."
)
figure(os.path.join(DEMO_FIG, "custom_cnn_confusion_matrix.png"),
      "Figure 3a: Custom CNN confusion matrix, AD vs CN (n=75).")
figure(os.path.join(DEMO_FIG, "custom_cnn_training_curves.png"),
      "Figure 3b: Custom CNN training curves, AD vs CN.")
figure(os.path.join(DEMO_FIG, "mobilenetv2_confusion_matrix.png"),
      "Figure 3c: MobileNetV2 confusion matrix, AD vs CN (n=75).")
figure(os.path.join(DEMO_FIG, "mobilenetv2_training_curves.png"),
      "Figure 3d: MobileNetV2 training curves, AD vs CN.")
figure(os.path.join(DEMO_FIG, "efficientnet_b0_confusion_matrix.png"),
      "Figure 3e: EfficientNet-B0 confusion matrix, AD vs CN (n=75).")
figure(os.path.join(DEMO_FIG, "efficientnet_b0_training_curves.png"),
      "Figure 3f: EfficientNet-B0 training curves, AD vs CN.")
para(
    "All three follow roughly the same pattern: validation accuracy tracks "
    "training accuracy pretty closely for the first several epochs before the gap "
    "starts to widen, and the checkpoint I actually keep (picked by validation "
    "loss) sits well before that gap gets big. That's a good sign -- it means the "
    "early-stopping setup is doing its job instead of just letting the models "
    "memorize the training set."
)
figure(os.path.join(FIG, "gradcam_mobilenetv2_ADvsCN.png"),
      "Figure 4: Grad-CAM attention for the AD-vs-CN model, showing attention "
      "concentrated around the ventricles and medial temporal lobe.")

h2("4.2 Four-Stage Classification (CN / EMCI / LMCI / AD) -- Still Exploratory")
para(
    "I also tried the harder four-stage version of the task, restricted to the "
    "ADNI-GO/2 cohort (618 subjects) so scanner era can't separate the classes on "
    "its own. Single-split results (93 test subjects, majority baseline 36.7%) "
    "are below -- none of these should be read as a settled result by "
    "themselves, given how wide the confidence interval is at this sample size."
)
table(["Model", "Subject accuracy", "Macro F1", "Soft-vote accuracy"],
      [["Custom CNN", "41.9%", "0.397", "40.9%"],
       ["MobileNetV2", "38.7%", "0.383", "36.6%"],
       ["EfficientNet-B0", "39.8%", "0.245", "37.6%"]])
para(
    "5-fold cross-validation over all 618 subjects gives the custom CNN 43.0% "
    "accuracy, 95% CI [39.2%, 47.0%], against a 36.7% baseline -- the bottom of "
    "that interval just clears the baseline, and it's the only version of this "
    "task that comes out statistically significant. But the confound check in "
    "Section 3.3 found that geometry alone accounts for about 4.2 of those 6.3 "
    "points, so I'd call this result suggestive at best, not established. There's "
    "also a labelling problem underneath the technical one: ADNI splits early "
    "from late mild cognitive impairment using a memory-test score cutoff, not "
    "anything visible on a scan, and ADNI's own documentation says that "
    "distinction isn't even kept consistent at follow-up visits. If I merge EMCI "
    "and LMCI into one class, accuracy jumps by about 16 points, and even that "
    "merged task doesn't reliably beat its own baseline. I'm not presenting this "
    "task as a finished result."
)
para(
    "Confusion matrices and training curves for all three models on this task "
    "are below. They all show the same basic failure mode: CN and AD, the two "
    "classes that are clinically furthest apart, get separated reasonably well, "
    "while EMCI and LMCI get mixed up heavily with each other and their "
    "neighbours -- which lines up visually with the labelling problem described "
    "above, separately from the confound numbers."
)
figure(os.path.join(FIG, "fourway_confusion_custom_cnn.png"),
      "Figure 5a: Custom CNN confusion matrix, four-stage task (n=93).")
figure(os.path.join(FIG, "fourway_curves_custom_cnn.png"),
      "Figure 5b: Custom CNN training curves, four-stage task.")
figure(os.path.join(FIG, "fourway_confusion_mobilenetv2.png"),
      "Figure 5c: MobileNetV2 confusion matrix, four-stage task (n=93).")
figure(os.path.join(FIG, "fourway_curves_mobilenetv2.png"),
      "Figure 5d: MobileNetV2 training curves, four-stage task.")
figure(os.path.join(FIG, "fourway_confusion_efficientnet_b0.png"),
      "Figure 5e: EfficientNet-B0 confusion matrix, four-stage task (n=93).")
figure(os.path.join(FIG, "fourway_curves_efficientnet_b0.png"),
      "Figure 5f: EfficientNet-B0 training curves, four-stage task.")
para(
    "EfficientNet-B0's confusion matrix makes the problem obvious: EMCI recall "
    "looks great (29 of 34), but that's mostly because everything else is getting "
    "dumped into that class too -- 23 of 24 CN subjects, 14 of 19 LMCI subjects, "
    "and 9 of 16 AD subjects all got predicted as EMCI. It basically learned to "
    "guess the biggest class instead of telling four stages apart, which is "
    "exactly what you'd expect once the label boundary itself is shaky, not "
    "really a sign that the architecture is bad."
)

h2("4.3 One Preprocessing Change That Actually Held Up Under Cross-Validation")
para(
    "I tried zeroing out the skull, scalp, and background around the brain "
    "(skull-stripping) as a possible improvement for AD-vs-CN, and it's the "
    "first preprocessing change in this project that survived paired 5-fold "
    "cross-validation: AUC went from 0.7643 to 0.8140, a paired improvement of "
    "+0.050 with a 95% CI of [+0.019, +0.082] that doesn't cross zero. The gain "
    "also held up in both directions of a cross-scanner-generation test (train on "
    "one ADNI phase, test on the other), which is evidence against this just "
    "being a new silhouette-based shortcut. This is the best-supported "
    "configuration I've measured so far, but it wasn't in the deployed model as "
    "of writing this."
)

h2("4.4 Discussion")
para(
    "Two things stood out across both tasks. First, bigger models aren't "
    "automatically better here -- when I deliberately switched to a slice-level "
    "split instead of a subject-level one (basically forcing leakage back in on "
    "purpose, as a control), MobileNetV2 and EfficientNet-B0 both jumped to "
    "92-96% accuracy versus roughly 55% honest, a much bigger leakage effect than "
    "the small custom CNN showed under the exact same manipulation. It seems like "
    "extra capacity that doesn't help detect atrophy is very good at memorizing "
    "individual people instead. Second, single train/test splits just aren't "
    "reliable at this dataset size -- I ran into the same lucky-split pattern "
    "three separate times over the course of this project, which is exactly why "
    "every headline number in this paper only gets reported after checking it "
    "under cross-validation."
)

# -------------------------------------------------------------- CONCLUSION -
h1("5. CONCLUSION")
para(
    "I built an Alzheimer's MRI classifier from raw ADNI DICOM data and treated "
    "the risk of the model taking shortcuts as a core part of the project, not "
    "an afterthought. Out of four confounds I tested for, I fixed three and I'm "
    "reporting the fourth as an open limitation. Under 5-fold subject-level "
    "cross-validation, the AD-vs-CN classifier reaches 74.1% accuracy and ROC "
    "AUC 0.784 [0.743, 0.826], which matches properly subject-split literature "
    "and a large MRI foundation model tested the same way. I also tried the "
    "harder four-stage task, but it's still tangled up with a geometry shortcut "
    "at the sample size I have, so I'm reporting it as unresolved rather than "
    "calling it done. If there's one takeaway from this project, it's that how "
    "carefully you evaluate a model -- splitting by subject, checking for "
    "confounds directly, reporting cross-validated numbers with confidence "
    "intervals -- changed my results more than which architecture I picked did."
)

# ------------------------------------------------------------- FUTURE WORK -
h1("6. FUTURE WORK")
bullet("Try an actual site-harmonization method (ComBat, CovBat, or DeepComBat) "
       "on the +8.0% site-identity confound instead of just reporting it.")
bullet("Take the geometry fix (isotropic resampling) I already built and run a "
       "full retrain and cross-validation pass with it, to see whether the "
       "four-stage task's margin survives once geometry is fully neutralized.")
bullet("Move from scoring slices independently to a real 3D convolutional model "
       "that looks at a person's whole slice stack at once. I've started this: "
       "5-fold cross-validation of a compact 3D CNN on AD-vs-CN is done as of "
       "writing this, and it came out at 64.3% accuracy [60.0, 68.3] and AUC "
       "0.7208 [0.6750, 0.7666] -- below the 2D headline of 74.1%/0.7845, so the "
       "3D approach didn't pay off on this first attempt. One fold looked like "
       "an early-stopping failure rather than a real result, so this is worth "
       "another pass before drawing a firm conclusion. I tried this in the first "
       "place because the 32 per-slice predictions used in Sections 4.1-4.2 turn "
       "out to be only about 1.3 independent measurements once you account for "
       "how correlated they are with each other -- a ceiling that only a model "
       "with real cross-slice context could get past.")
bullet("Bring in longitudinal scans, since how fast the hippocampus shrinks "
       "between visits is a stronger predictor of MCI-to-AD conversion than any "
       "single scan used here.")
bullet("Deploy the skull-stripped setup from Section 4.3 once its cross-era "
       "robustness gets a bit more confirmation.")

# -------------------------------------------------------------- REFERENCES -
h1("7. REFERENCES")
refs = [
    "[1] Mujahid, M., Rehman, A., Alam, T., Alamri, F. S., Fati, S. M., & Saba, T. "
    "\"An Efficient Ensemble Approach for Alzheimer's Disease Detection Using "
    "Adaptive Synthetic Technique and Deep Learning,\" Diagnostics, 2023.",
    "[2] Ali, M. U., Hussain, S. J., Khalid, M., Farrash, M., Lahza, H. F. M., & "
    "Zafar, A. \"MRI-Driven Alzheimer's Disease Diagnosis Using Deep Network "
    "Fusion and Optimal Feature Selection,\" Bioengineering, 2024.",
    "[3] Tan, M., & Le, Q. \"EfficientNet: Rethinking Model Scaling for "
    "Convolutional Neural Networks,\" ICML, 2019.",
    "[4] Howard, A. et al. \"MobileNets: Efficient Convolutional Neural Networks "
    "for Mobile Vision Applications,\" 2017.",
    "[5] Krizhevsky, A., Sutskever, I., & Hinton, G. \"ImageNet Classification "
    "with Deep Convolutional Neural Networks,\" NeurIPS, 2012.",
    "[6] Litjens, G. et al. \"A Survey on Deep Learning in Medical Image "
    "Analysis,\" Medical Image Analysis, 2017.",
    "[7] Goodfellow, I., Bengio, Y., & Courville, A. Deep Learning. MIT Press, "
    "2016.",
    "[8] Ansart, M. et al. \"Cross-Cohort Generalizability of Deep and "
    "Conventional Machine Learning for MRI-based Diagnosis and Prediction of "
    "Alzheimer's Disease,\" arXiv:2012.08769.",
    "[9] Tinauer, C. et al. \"Skull-Stripping Induces Shortcut Learning in "
    "Alzheimer's Disease MRI Classification,\" arXiv:2501.15831.",
    "[10] \"3D MRI-Based Alzheimer's Disease Classification Using Multi-Modal "
    "3D CNN with Leakage-Aware Subject-Level Evaluation,\" arXiv:2603.17304.",
    "[11] Wu, M. et al. \"BrainDINO: A Brain MRI Foundation Model for "
    "Generalizable Clinical Representation Learning,\" arXiv:2604.27277, 2026.",
]
for r in refs:
    doc.add_paragraph(r)

doc.save(OUT)
print(f"saved: {OUT}")
