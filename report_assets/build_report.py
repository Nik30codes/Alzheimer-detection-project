"""Builds the ETE research paper as a .docx, matching the required section structure,
using only real numbers pulled from this project's own result files and real figures
generated from actual model runs. Run once: python report_assets/build_report.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Nikunj_Bhalla_2430030052_ETE_Research_Paper.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(htext)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")
    return t


def figure(path, caption, width=4.6):
    if not os.path.exists(path):
        para(f"[figure missing: {os.path.basename(path)}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)


FIG = os.path.join(HERE)
DEMO_FIG = os.path.join(HERE, "..", "demo", "results")

# ---------------------------------------------------------------- TITLE ----
title = doc.add_heading(
    "Confound-Aware Deep Learning Classification of Alzheimer's Disease "
    "from Real ADNI MRI Scans", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Nikunj Bhalla")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -------------------------------------------------------------- ABSTRACT ---
h1("ABSTRACT")
para(
    "Automated staging of Alzheimer's disease from structural MRI is attractive "
    "because atrophy patterns are visible well before the disease reaches its "
    "clinical end stage, but much of the published deep-learning literature on this "
    "task reports 80-99% accuracy under evaluation conditions that let information "
    "about a subject or a scanning protocol leak from training into test data. We "
    "build a classification pipeline from raw ADNI DICOM data -- 853 subjects across "
    "two scanner generations, not a pre-packaged image set -- and check for four "
    "specific sources of such leakage before reporting any result: subject-level "
    "leakage, a scanner-era/diagnosis confound, a misaligned slice-extraction band, "
    "and an acquisition-geometry artifact. Three are corrected; the fourth, imaging "
    "site, is quantified and left as an open limitation. A compact custom CNN, "
    "MobileNetV2, and EfficientNet-B0 are compared, all trained from scratch after "
    "ImageNet pretraining was found to hurt rather than help on this modality. For "
    "Alzheimer's-versus-cognitively-normal classification, 5-fold subject-level "
    "cross-validation over 501 subjects gives 74.1% accuracy and ROC AUC 0.784 "
    "(95% CI [0.743, 0.826]), consistent with properly subject-split literature and "
    "with a recent large-scale MRI foundation model evaluated on the same task. A "
    "harder four-stage task (cognitively normal / early and late mild cognitive "
    "impairment / Alzheimer's) is also attempted; it is not yet distinguishable from "
    "a geometry-driven shortcut at the available sample size and is reported as "
    "unresolved rather than as a finished result."
)

# ----------------------------------------------------------- INTRODUCTION --
h1("1. INTRODUCTION")
para(
    "Alzheimer's disease is the most common cause of dementia. Its course is a "
    "gradual loss of memory and independent function over years, and structural MRI "
    "is central to its clinical workup because atrophy in the hippocampus and "
    "surrounding medial temporal structures becomes visible well before end-stage "
    "disease. A classifier that reads this signal reliably could support earlier, "
    "more consistent screening than manual radiological review alone."
)
para(
    "The practical difficulty in this field is not training a convolutional network "
    "that scores well on a held-out test set -- it is establishing that it learned "
    "atrophy rather than a shortcut. Two properties of MRI classification datasets "
    "make shortcuts easy to acquire by accident. A single subject typically "
    "contributes dozens of near-duplicate slices, so any split not enforced at the "
    "subject level lets a network partly memorize individuals rather than generalize "
    "across them. Acquisition protocol -- scanner model, coil, voxel spacing, field "
    "of view -- also tends to correlate with diagnosis for reasons that have nothing "
    "to do with anatomy, simply because different cohorts were scanned on different "
    "equipment at different times. A model can reach a high score by reading either "
    "of these instead of the brain."
)
para(
    "We treat that risk as a central methodological question rather than an "
    "afterthought. Starting from 853 subjects' worth of raw ADNI DICOM series, the "
    "pipeline reconstructs a 3D volume per subject, reslices it into the standard "
    "axial view, and extracts a millimetre-anchored slice band before any "
    "train/test split is drawn. Four distinct potential shortcuts are then searched "
    "for, each with its own measurement, and the three found to be real are "
    "corrected before any headline number is reported. This paper's contributions:"
)
bullet("A 3D-volume-aware, millimetre-anchored slice-extraction pipeline built "
       "directly from raw sagittal DICOM series, rather than pre-extracted images.")
bullet("A systematic, quantified search for four separate sources of shortcut "
       "learning -- subject leakage, scanner-era confounding, anatomical "
       "misalignment, and acquisition-geometry artifacts -- each measured with its "
       "own method rather than assumed absent.")
bullet("A controlled comparison of a custom CNN, MobileNetV2, and EfficientNet-B0, "
       "all trained from scratch, with the empirical finding that ImageNet "
       "pretraining is a net negative for grayscale MRI classification in this "
       "setting.")
bullet("A statistically validated headline result -- 5-fold subject-level "
       "cross-validation with confidence intervals, not a single train/test split "
       "-- after demonstrating that single-split estimates in this exact pipeline "
       "can overstate accuracy by more than ten points.")
bullet("An honest accounting of what remains unresolved, including a four-stage "
       "classification result that is explicitly not claimed as established.")

# -------------------------------------------------------- LITERATURE REVIEW-
h1("2. LITERATURE REVIEW")
para(
    "Convolutional networks are the dominant tool for MRI-based Alzheimer's "
    "classification, learning discriminative features directly from pixel data "
    "without hand-engineered morphometric measurements. Transfer learning from "
    "ImageNet-pretrained backbones such as EfficientNet [3] and MobileNet [4] is "
    "common in this literature, on the assumption that low-level convolutional "
    "features learned from natural images transfer usefully when target-domain "
    "data is limited. As shown in Section 4, we find this assumption does not hold "
    "for grayscale MRI."
)
para(
    "A second, less frequently addressed strand of the literature concerns "
    "evaluation validity rather than architecture. Ansart et al. [8] find that "
    "models validated within a single site or cohort routinely fail to reproduce "
    "their reported accuracy on external data, implicating cohort-specific "
    "shortcuts rather than genuine anatomical signal. Tinauer et al. [9] show that "
    "skull-stripping, often treated as a neutral preprocessing step, can itself "
    "introduce shortcut cues depending on the dataset -- a Clever Hans effect that "
    "argues for validating preprocessing choices empirically rather than assuming "
    "them safe. A more recent leakage-aware 3D CNN study [10] makes subject-level "
    "partitioning an explicit, reported constraint precisely because so much prior "
    "work leaves it unstated. BrainDINO [11], a self-distilled foundation model "
    "pretrained on 6.6 million MRI slices from twenty datasets, reports "
    "subject-disjoint AD-vs-CN performance of AUC 0.850 (95% CI [0.754, 0.947]) -- "
    "a useful, methodologically careful reference point for what this task's "
    "ceiling looks like without shortcut inflation."
)
para(
    "Site and scanner harmonization is a related, named problem in the wider "
    "neuroimaging literature: tools such as ComBat and its extensions (CovBat, "
    "DeepComBat) were developed specifically to remove scanner- and "
    "site-attributable variance from imaging features before downstream modelling. "
    "Our own site-effect measurement (Section 4.1) is consistent with that "
    "literature's premise; applying one of these harmonization methods is left as "
    "future work rather than attempted here."
)
para(
    "This work sits at the intersection of the two strands above: standard CNN "
    "architectures from the first, evaluated under the methodological discipline "
    "argued for by the second -- explicit confound measurement, subject-level "
    "splitting, and cross-validated reporting -- rather than a single held-out test "
    "accuracy taken at face value."
)

# ------------------------------------------------------------- METHODOLOGY -
h1("3. METHODOLOGY")

h2("3.1 Dataset Description")
para(
    "Data were obtained directly from the ADNI (Alzheimer's Disease Neuroimaging "
    "Initiative) archive as raw DICOM series, downloaded in two batches spanning "
    "ADNI's earlier (~2005-2007) and later (2011+) acquisition phases -- referred "
    "to below as ADNI1 and ADNI-GO/2. The combined pool contains 853 subjects "
    "across four diagnostic categories (Table 1):"
)
table(["Class", "ADNI1", "ADNI-GO/2", "Total"],
      [["Cognitively Normal (CN)", 127, 158, 285],
       ["Alzheimer's Disease (AD)", 108, 108, 216],
       ["Early MCI (EMCI)", 0, 227, 227],
       ["Late MCI (LMCI)", 0, 125, 125]])
para(
    "Each subject contributes one T1-weighted MPRAGE-family series. Subjects were "
    "split 70/15/15 into train/validation/test sets at the subject level, before "
    "any slice image was generated, so that no person's scans appear in more than "
    "one split."
)
figure(os.path.join(FIG, "class_distribution.png"),
      "Figure 1: Class distribution across the dataset.")

h2("3.2 Data Preprocessing")
para(
    "Each subject's sagittally-acquired DICOM series is stacked into a 3D volume "
    "and reconstructed into the standard axial (top-down) orientation. A 32-slice "
    "band is then extracted, anchored not as a fixed fraction of image height but "
    "as a fixed physical depth (48-92mm) below the automatically detected vertex of "
    "the skull, using each scan's own pixel spacing. This anchoring was adopted "
    "after an earlier fraction-based approach was found to drift onto different "
    "anatomy in different subjects -- in some cases missing the hippocampus "
    "entirely, the single most diagnostically relevant structure for this task. "
    "Reslicing across independently-acquired sagittal images introduces "
    "reconstruction-grain noise, which is removed with Non-Local Means denoising. "
    "Because ADNI1 scans are natively lower resolution than ADNI-GO/2 scans, every "
    "image is routed through a common intermediate resolution before final resizing "
    "to 224x224, so that resolution itself does not become a proxy for scanner "
    "generation (and therefore, given the class distribution above, a proxy for "
    "diagnosis)."
)
figure(os.path.join(FIG, "qc_v3_range.png"),
      "Figure 2: Example extracted axial slice range across the anchored band, "
      "confirming consistent anatomical coverage.")

h2("3.3 Confound Detection Methodology")
para(
    "Before any classification result is reported, each candidate confound is "
    "tested with the same method: a simple classifier (majority-baseline comparison, "
    "or a shallow model trained on non-image metadata alone -- scanner era, native "
    "resolution, voxel spacing) is asked to predict the diagnostic label using only "
    "that one piece of information, with no image pixels involved. If it scores "
    "above the majority baseline, that variable is a usable shortcut and any "
    "image-based model could in principle be exploiting it rather than anatomy. "
    "Four such tests were run:"
)
table(["#", "Candidate confound", "Test result", "Status"],
      [["1", "Same subject's slices split across train/test",
        "+36.9 accuracy points when deliberately reintroduced", "Fixed: subject-wise split"],
       ["2", "Scanner era (ADNI1 vs GO/2) correlated with class",
        "Era predictable at 95-100% pre-fix; diagnosis near chance", "Fixed: era-balanced expansion"],
       ["3", "Slice band anatomical alignment",
        "Band drifted onto different anatomy per subject; some missed hippocampus", "Fixed: mm-anchored band"],
       ["4", "Acquisition geometry / anisotropy",
        "Metadata alone: +4.2% on 4-way task; -4.6% (clean) on AD-vs-CN", "Partially fixed (isotropic resampling)"]])
para(
    "A fifth variable, imaging site, was also measured and found to carry a "
    "non-trivial signal (+8.0% over baseline on the AD-vs-CN task from site "
    "identity alone). This estimate is likely inflated by the large number of "
    "site levels relative to sample size, and per-era test accuracy for the final "
    "model is balanced (evidence against pure protocol reading), but the effect is "
    "not fully resolved and is reported here as an open limitation rather than "
    "corrected."
)

h2("3.4 Model Architectures")
para(
    "Three architectures were compared, all consuming a single 224x224 axial slice "
    "per forward pass:"
)
bullet("Custom CNN (~1.2M parameters): four convolutional blocks (Conv-BatchNorm-"
       "ReLU x2, max-pool) with channel widths 32-64-128-256, followed by global "
       "average pooling and a two-layer dropout-regularized classification head. "
       "Global average pooling was used in place of a flatten-and-dense head "
       "specifically to control parameter count on a dataset of this size.")
bullet("MobileNetV2 (~2.2M parameters): torchvision implementation, inverted "
       "residual blocks with depthwise-separable convolutions.")
bullet("EfficientNet-B0 (~4.0M parameters): torchvision implementation, compound "
       "depth/width/resolution scaling.")
para(
    "ImageNet-pretrained initialization was tested for MobileNetV2 under three "
    "fine-tuning strategies (full unfreeze, partial unfreeze with frozen "
    "BatchNorm, partial unfreeze with adaptive BatchNorm) and found to underperform "
    "the same architecture trained from random initialization by 15-20 accuracy "
    "points in every configuration. This is consistent with the pretrained "
    "features encoding natural-photograph statistics that do not transfer well to "
    "grayscale MRI duplicated into synthetic RGB channels. All results reported "
    "below therefore use random initialization (training from scratch) for every "
    "architecture."
)

h2("3.5 Training Setup")
table(["Parameter", "Value"],
      [["Input size", "224 x 224, 1 or 3 channels"],
       ["Optimizer", "Adam, lr 1e-3"],
       ["Weight decay", "1e-4"],
       ["Batch size", "32"],
       ["Loss", "Class-weighted cross-entropy"],
       ["LR schedule", "ReduceLROnPlateau (factor 0.5, patience 2)"],
       ["Early stopping", "Patience 7 epochs, best checkpoint restored"],
       ["Max epochs", "40"],
       ["Augmentation", "Random affine (rotate/translate/scale) + horizontal flip"],
       ["Precision", "Mixed precision (fp16 autocast + gradient scaling)"]])

h2("3.6 Evaluation Protocol")
para(
    "Each subject's slice-level predictions are aggregated by averaging softmax "
    "probabilities across their slices (soft voting) into a single subject-level "
    "prediction, since the clinical unit of interest is the person, not an "
    "individual 2D image. The primary evaluation for the binary AD-vs-CN task is "
    "5-fold subject-level cross-validation: every subject receives exactly one "
    "out-of-fold prediction, and the pooled result is reported with a 95% Wilson "
    "confidence interval for accuracy and a Hanley-McNeil confidence interval for "
    "ROC AUC. This was adopted after observing that a single 75-subject test split "
    "of this exact pipeline reported 82.7% accuracy / AUC 0.906, while 5-fold "
    "cross-validation of the identical configuration over all 501 subjects gave "
    "74.1% / AUC 0.784 with non-overlapping confidence intervals -- i.e. the single "
    "split was a lucky sample, not a reproducible estimate. The decision threshold "
    "for converting a probability into a class label is fit on validation subjects "
    "via Youden's J statistic and never on the test fold."
)

# ------------------------------------------------------ RESULTS/DISCUSSION -
h1("4. RESULTS AND DISCUSSION")

h2("4.1 AD vs Cognitively Normal -- Primary Result")
para(
    "Table 3 reports each architecture's performance on a single held-out test "
    "split (75 subjects) for illustration, alongside the statistically validated "
    "5-fold cross-validated result. Only the cross-validated figure should be read "
    "as the project's headline, for the reason given in Section 3.6."
)
table(["Model", "Split", "n", "Accuracy", "Macro F1", "ROC AUC"],
      [["Custom CNN", "single split", 75, "73.3%", "0.725", "0.838"],
       ["MobileNetV2", "single split", 75, "76.0%", "0.758", "0.852"],
       ["EfficientNet-B0", "single split", 75, "78.7%", "0.786", "0.876"],
       ["MobileNetV2", "5-fold CV (headline)", 501, "74.1% [70.0, 77.7]", "--",
        "0.784 [0.743, 0.826]"]])
para(
    "The cross-validated AUC interval excludes 0.5 by a wide margin, indicating a "
    "genuine, non-chance ranking signal. It is also consistent with the wider "
    "literature: properly subject-split studies on comparably sized ADNI cohorts "
    "report binary AD-vs-CN accuracy under 71% [reviewed informally against 9], and "
    "BrainDINO, a foundation model pretrained on 6.6 million MRI slices, reports "
    "AUC 0.850 [0.754, 0.947] [11] on the same comparison -- an interval that "
    "overlaps this study's result despite this study using two orders of magnitude "
    "less pretraining data and no foundation-model backbone."
)
para(
    "Per-model confusion matrices and training curves for all three architectures "
    "follow, each from the same 75-subject held-out split as Table 3, so the visual "
    "pattern of errors and the shape of the learning curve can be compared directly "
    "across architectures."
)
figure(os.path.join(DEMO_FIG, "custom_cnn_confusion_matrix.png"),
      "Figure 3a: Custom CNN confusion matrix, AD vs CN (n=75).")
figure(os.path.join(DEMO_FIG, "custom_cnn_training_curves.png"),
      "Figure 3b: Custom CNN training curves, AD vs CN.")
figure(os.path.join(DEMO_FIG, "mobilenetv2_confusion_matrix.png"),
      "Figure 3c: MobileNetV2 confusion matrix, AD vs CN (n=75).")
figure(os.path.join(DEMO_FIG, "mobilenetv2_training_curves.png"),
      "Figure 3d: MobileNetV2 training curves, AD vs CN.")
figure(os.path.join(DEMO_FIG, "efficientnet_b0_confusion_matrix.png"),
      "Figure 3e: EfficientNet-B0 confusion matrix, AD vs CN (n=75).")
figure(os.path.join(DEMO_FIG, "efficientnet_b0_training_curves.png"),
      "Figure 3f: EfficientNet-B0 training curves, AD vs CN.")
para(
    "All three show the same qualitative pattern: validation accuracy tracks "
    "training accuracy reasonably closely for the first several epochs before the "
    "gap widens, and the best checkpoint (selected on validation loss) sits well "
    "before the point where that gap becomes large -- consistent with the "
    "early-stopping and checkpoint-selection choices in Section 3.5 doing their "
    "intended job rather than the models simply memorising the training set."
)
figure(os.path.join(FIG, "gradcam_mobilenetv2_ADvsCN.png"),
      "Figure 4: Grad-CAM attention overlays for the AD-vs-CN model, showing "
      "attention concentrated on ventricular and medial temporal regions.")

h2("4.2 Four-Stage Classification (CN / EMCI / LMCI / AD) -- Exploratory")
para(
    "The harder four-stage task was also attempted, restricted to the ADNI-GO/2 "
    "cohort (618 subjects) so that scanner era does not separate the classes. "
    "Single-split subject-level results (93 test subjects, majority baseline "
    "36.7%) are shown below; none should be read as an established result on "
    "their own, given the confidence-interval width at this sample size."
)
table(["Model", "Subject accuracy", "Macro F1", "Soft-vote accuracy"],
      [["Custom CNN", "41.9%", "0.397", "40.9%"],
       ["MobileNetV2", "38.7%", "0.383", "36.6%"],
       ["EfficientNet-B0", "39.8%", "0.245", "37.6%"]])
para(
    "5-fold cross-validation over all 618 subjects gives a pooled estimate for the "
    "custom CNN of 43.0% accuracy, 95% CI [39.2%, 47.0%] against the 36.7% "
    "baseline -- the interval's lower bound clears the baseline, which is the only "
    "arm of this task found to be statistically significant. However, the "
    "confound audit in Section 4.1 (Table with the four tested confounds) found "
    "that acquisition geometry alone supplies roughly 4.2 of this 6.3-point "
    "margin, so this result is reported as suggestive rather than established. It "
    "is compounded by a labelling issue rather than a purely technical one: ADNI "
    "distinguishes early from late mild cognitive impairment by a memory-test "
    "score cutoff rather than an imaging finding, and its own documentation notes "
    "that the distinction is not consistently maintained at follow-up visits. "
    "Consistent with this, collapsing EMCI and LMCI into one class recovers "
    "roughly 16 accuracy points, and even the merged task does not clear its own "
    "baseline reliably. This task is deliberately not presented as a finished "
    "result."
)
para(
    "Confusion matrices and training curves for all three architectures on this "
    "task follow. The same qualitative failure mode appears in every one: CN and "
    "AD (the two classes furthest apart clinically) are separated reasonably well, "
    "while EMCI and LMCI are heavily confused with each other and with their "
    "clinical neighbours -- visual confirmation of the labelling-boundary problem "
    "described above, independent of the confound-audit numbers."
)
figure(os.path.join(FIG, "fourway_confusion_custom_cnn.png"),
      "Figure 5a: Custom CNN confusion matrix, four-stage task (n=93).")
figure(os.path.join(FIG, "fourway_curves_custom_cnn.png"),
      "Figure 5b: Custom CNN training curves, four-stage task.")
figure(os.path.join(FIG, "fourway_confusion_mobilenetv2.png"),
      "Figure 5c: MobileNetV2 confusion matrix, four-stage task (n=93).")
figure(os.path.join(FIG, "fourway_curves_mobilenetv2.png"),
      "Figure 5d: MobileNetV2 training curves, four-stage task.")
figure(os.path.join(FIG, "fourway_confusion_efficientnet_b0.png"),
      "Figure 5e: EfficientNet-B0 confusion matrix, four-stage task (n=93).")
figure(os.path.join(FIG, "fourway_curves_efficientnet_b0.png"),
      "Figure 5f: EfficientNet-B0 training curves, four-stage task.")
para(
    "EfficientNet-B0's confusion matrix is the clearest illustration: EMCI recall "
    "is high (29 of 34), but that comes largely at the expense of the other three "
    "classes collapsing into it -- 23 of 24 CN subjects, 14 of 19 LMCI subjects, "
    "and 9 of 16 AD subjects are all predicted EMCI. The model has effectively "
    "learned to guess the largest class rather than distinguish four stages, "
    "which is exactly the failure mode the merged-class analysis above predicts "
    "once the label boundary is unreliable, rather than a purely architectural "
    "weakness."
)

h2("4.3 A Preprocessing Change That Survived Cross-Validation")
para(
    "Zeroing skull, scalp, and background pixels (skull-stripping) around the "
    "brain, tested as a candidate preprocessing improvement for AD-vs-CN, produced "
    "the first preprocessing change in this project to survive paired 5-fold "
    "cross-validation: AUC improved from 0.7643 to 0.8140 (paired delta +0.050, "
    "95% CI [+0.019, +0.082], excluding zero), and the gain held in both "
    "directions of a cross-scanner-generation transfer test (train on one ADNI "
    "phase, test on the other), which argues against the gain being a new "
    "silhouette-based shortcut rather than genuine attention improvement. This "
    "configuration is the best-supported one measured in this project but was not "
    "yet in the deployed model at the time of writing."
)

h2("4.4 Discussion")
para(
    "Two patterns recur across both tasks. First, model capacity does not "
    "translate simply into better generalization on this dataset: when a slice-"
    "level split was deliberately used instead of a subject-level one (a "
    "leakage-injection control), the larger MobileNetV2 and EfficientNet-B0 "
    "reached 92-96% accuracy versus roughly 55% honest, a substantially larger "
    "leakage effect than the small custom CNN showed under the same manipulation "
    "-- capacity that is not useful for detecting atrophy is apparently very "
    "useful for memorizing individual subjects. Second, single train/test splits "
    "on a dataset of this size are not reliable enough to report as final numbers: "
    "this project reproduced the same lucky-split pattern three separate times "
    "across different experiments, and every headline figure in this paper is "
    "reported only after confirming it under cross-validation for that reason."
)

# -------------------------------------------------------------- CONCLUSION -
h1("5. CONCLUSION")
para(
    "We built an Alzheimer's MRI classification pipeline from raw ADNI DICOM data "
    "and treated the risk of shortcut learning as a first-class methodological "
    "question rather than an assumption. Of four candidate "
    "confounds tested, three were corrected and one is reported as an open "
    "limitation. Under 5-fold subject-level cross-validation, the resulting "
    "AD-vs-CN classifier reaches 74.1% accuracy and ROC AUC 0.784 [0.743, 0.826], "
    "a result consistent with properly subject-split literature and with a "
    "large-scale MRI foundation model evaluated on the same comparison. The "
    "harder four-stage classification task was also attempted and found not yet "
    "distinguishable from a geometry-driven shortcut at the available sample "
    "size, and is reported honestly as unresolved. The central conclusion of this "
    "work is methodological as much as architectural: on a dataset of this size "
    "and structure, evaluation discipline -- subject-wise splitting, explicit "
    "confound measurement, and cross-validated reporting with confidence "
    "intervals -- changes the reported result by more than architecture choice "
    "does."
)

# ------------------------------------------------------------- FUTURE WORK -
h1("6. FUTURE WORK")
bullet("Apply a named site-harmonization method (ComBat, CovBat, or DeepComBat) "
       "to address the measured +8.0% site-identity confound directly, rather "
       "than only reporting it.")
bullet("Extend the acquisition-geometry fix (isotropic resampling) already built "
       "for this pipeline into a full retraining and cross-validation pass, to "
       "confirm whether the four-stage task's margin survives once geometry is "
       "fully neutralized.")
bullet("Replace independent 2D slice scoring with a true 3D convolutional model "
       "that consumes a subject's full slice stack jointly. Preliminary "
       "cross-validation of a compact 3D CNN on the AD-vs-CN task is in progress "
       "at the time of writing (2 of 5 folds complete: 53.5% and 73.0% subject "
       "accuracy respectively), motivated by a finding that the 32 per-slice "
       "predictions used in Sections 4.1-4.2 are only about 1.3 effectively "
       "independent measurements once their correlation is accounted for -- a "
       "ceiling that only an architecture with genuine cross-slice context can "
       "exceed.")
bullet("Incorporate longitudinal scans, since the rate of hippocampal atrophy "
       "between visits is a stronger reported predictor of MCI-to-AD conversion "
       "than any single-timepoint measurement used here.")
bullet("Deploy the skull-stripped configuration from Section 4.3 to production "
       "once its cross-era robustness is further confirmed.")

# -------------------------------------------------------------- REFERENCES -
h1("7. REFERENCES")
refs = [
    "[1] Mujahid, M., Rehman, A., Alam, T., Alamri, F. S., Fati, S. M., & Saba, T. "
    "\"An Efficient Ensemble Approach for Alzheimer's Disease Detection Using "
    "Adaptive Synthetic Technique and Deep Learning,\" Diagnostics, 2023.",
    "[2] Ali, M. U., Hussain, S. J., Khalid, M., Farrash, M., Lahza, H. F. M., & "
    "Zafar, A. \"MRI-Driven Alzheimer's Disease Diagnosis Using Deep Network "
    "Fusion and Optimal Feature Selection,\" Bioengineering, 2024.",
    "[3] Tan, M., & Le, Q. \"EfficientNet: Rethinking Model Scaling for "
    "Convolutional Neural Networks,\" ICML, 2019.",
    "[4] Howard, A. et al. \"MobileNets: Efficient Convolutional Neural Networks "
    "for Mobile Vision Applications,\" 2017.",
    "[5] Krizhevsky, A., Sutskever, I., & Hinton, G. \"ImageNet Classification "
    "with Deep Convolutional Neural Networks,\" NeurIPS, 2012.",
    "[6] Litjens, G. et al. \"A Survey on Deep Learning in Medical Image "
    "Analysis,\" Medical Image Analysis, 2017.",
    "[7] Goodfellow, I., Bengio, Y., & Courville, A. Deep Learning. MIT Press, "
    "2016.",
    "[8] Ansart, M. et al. \"Cross-Cohort Generalizability of Deep and "
    "Conventional Machine Learning for MRI-based Diagnosis and Prediction of "
    "Alzheimer's Disease,\" arXiv:2012.08769.",
    "[9] Tinauer, C. et al. \"Skull-Stripping Induces Shortcut Learning in "
    "Alzheimer's Disease MRI Classification,\" arXiv:2501.15831.",
    "[10] \"3D MRI-Based Alzheimer's Disease Classification Using Multi-Modal "
    "3D CNN with Leakage-Aware Subject-Level Evaluation,\" arXiv:2603.17304.",
    "[11] Wu, M. et al. \"BrainDINO: A Brain MRI Foundation Model for "
    "Generalizable Clinical Representation Learning,\" arXiv:2604.27277, 2026.",
]
for r in refs:
    doc.add_paragraph(r)

doc.save(OUT)
print(f"saved: {OUT}")
